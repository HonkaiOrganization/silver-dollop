import os
import io
import base64
import logging
import uuid
import re

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp")


def _extract_keyframes_from_b64(clip_b64: str, max_frames: int = 8) -> list[np.ndarray]:
    """从 base64 编码的视频中均匀抽取关键帧，返回 BGR 图像列表。"""
    data = base64.b64decode(clip_b64)
    os.makedirs(TEMP_DIR, exist_ok=True)
    tmp_path = os.path.join(TEMP_DIR, f"export_clip_{uuid.uuid4().hex[:8]}.mp4")
    try:
        with open(tmp_path, 'wb') as f:
            f.write(data)

        cap = cv2.VideoCapture(tmp_path)
        if not cap.isOpened():
            return []

        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        if total <= 0:
            cap.release()
            return []

        num = min(max_frames, total)
        step = max(1, total // num)
        indices = list(range(0, total, step))[:num]

        frames = []
        for idx in indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ret, frame = cap.read()
            if ret:
                frames.append(frame)
        cap.release()
        return frames
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _create_mosaic(frames: list[np.ndarray], cols: int = 4, target_width: int = 1200) -> np.ndarray | None:
    """将多帧图像拼成一张网格图。"""
    if not frames:
        return None

    rows = (len(frames) + cols - 1) // cols
    h, w = frames[0].shape[:2]
    cell_w = target_width // cols
    cell_h = int(h * cell_w / w)
    canvas_h = cell_h * rows
    canvas_w = cell_w * cols
    canvas = np.zeros((canvas_h, canvas_w, 3), dtype=np.uint8)

    for i, frame in enumerate(frames):
        r, c = divmod(i, cols)
        resized = cv2.resize(frame, (cell_w, cell_h), interpolation=cv2.INTER_AREA)
        canvas[r * cell_h:(r + 1) * cell_h, c * cell_w:(c + 1) * cell_w] = resized

    return canvas


def _add_markdown_text(doc, text: str, parent=None):
    """将 Markdown 文本简单解析后添加到文档中。处理加粗、列表、标题。"""
    from docx.oxml.ns import qn

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=4)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith('- **') or stripped.startswith('- '):
            content = stripped.lstrip('- ').strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, content)
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)


def _add_rich_text(paragraph, text: str):
    """解析 Markdown 加粗语法并添加到段落。"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def export_vlm_report(
    output_path: str,
    sections: list[dict],
    summary: str,
    infer_result: dict | None = None,
    csv_name: str = "",
):
    """
    将 VLM 分析结果导出为 Word 文档。

    Args:
        output_path: 输出 .docx 文件路径
        sections: VLM 分析段落列表，每项包含 title/prob/start_frame/end_frame/analysis/clip_b64/figure_number
        summary: 总结文本
        infer_result: 推理结果字典（可选，用于添加概要信息）
        csv_name: 分析的 CSV 文件名
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')

    title = doc.add_heading('跳绳动作分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if csv_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"分析文件：{csv_name}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    if infer_result:
        doc.add_heading('推理概要', level=1)
        r = infer_result.get('results', {})
        for fname, result in r.items():
            if result.get('status') != 'ok':
                continue
            label_cn = "正常" if result['predicted_label'] == 'normal' else "异常"
            doc.add_paragraph(f"判定结果：{label_cn}")
            doc.add_paragraph(f"置信度：{result['confidence']:.4f}")
            p_normal = result['probabilities']['normal']
            p_abnormal = result['probabilities']['abnormal']
            doc.add_paragraph(f"P(正常)：{p_normal:.4f} P(异常)：{p_abnormal:.4f}")
            doc.add_paragraph(f"滑动窗口数：{result['num_windows']} 总帧数：{result['num_frames']}")
            break

    if not sections:
        doc.add_paragraph("无分析数据。")
        doc.save(output_path)
        return

    doc.add_heading('问题分析', level=1)

    tmp_images = []

    for sec in sections:
        fig_num = sec.get('figure_number', 0)
        title_text = sec['title']
        prob = sec['prob']
        start_frame = sec['start_frame']
        end_frame = sec['end_frame']
        analysis = sec.get('analysis', '')
        clip_b64 = sec.get('clip_b64')

        doc.add_heading(title_text, level=2)

        p = doc.add_paragraph()
        run = p.add_run(f"帧范围：{start_frame} - {end_frame}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.add_run("    ")
        run2 = p.add_run(f"异常概率：{prob:.1%}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(128, 128, 128)

        if clip_b64:
            frames = _extract_keyframes_from_b64(clip_b64, max_frames=8)
            mosaic = _create_mosaic(frames, cols=4, target_width=1200)
            if mosaic is not None:
                os.makedirs(TEMP_DIR, exist_ok=True)
                img_path = os.path.join(TEMP_DIR, f"mosaic_{fig_num}_{uuid.uuid4().hex[:8]}.png")
                cv2.imwrite(img_path, mosaic)
                tmp_images.append(img_path)

                fig_caption = doc.add_paragraph()
                fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = fig_caption.add_run(f"图{fig_num} 问题片段关键帧截图")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)

                doc.add_picture(img_path, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if analysis:
            doc.add_paragraph()
            _add_markdown_text(doc, analysis)

        doc.add_paragraph()

    if summary:
        doc.add_heading('总结与改进优先级', level=1)
        _add_markdown_text(doc, summary)

    doc.save(output_path)

    for path in tmp_images:
        try:
            os.unlink(path)
        except OSError:
            pass

    logger.info("报告已导出: %s", output_path)
